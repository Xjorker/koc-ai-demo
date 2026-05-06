# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_para(doc, text, bold=False, italic=False, size=12):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return para

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# ========== 封面 ==========
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('KOC AI增长助手', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('腾讯PCG校园AI产品创意大赛')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

track = doc.add_paragraph('赛道⑤ - PCG技术线-智能营销/AI小游戏方向')
track.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in track.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
doc.add_paragraph()

intro = doc.add_paragraph('参赛说明文档')
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in intro.runs:
    run.bold = True
    run.font.size = Pt(18)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph('最后更新时间：2026年5月6日')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in info.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

info2 = doc.add_paragraph('文档版本：v2.0（提交版）')
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in info2.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ========== 一、产品概述 ==========
add_heading(doc, '一、产品概述', 1)
add_heading(doc, '1.1 参赛信息', 2)

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['项目', '内容']
data = [
    ['参赛赛道', '赛道⑤ - PCG技术线-智能营销/AI小游戏方向'],
    ['作品名称', 'KOC AI增长助手'],
    ['产品类型', '面向KOC的社媒AI Agent'],
    ['核心技术', '大语言模型（LLM）+ LangGraph Agent + 社交媒体数据分析'],
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = text

doc.add_paragraph()
add_heading(doc, '1.2 一句话介绍', 2)

quote = doc.add_paragraph()
quote_run = quote.add_run('KOC AI增长助手：让每个普通人都能轻松玩转社交媒体，通过AI提供账号诊断、智能选题、内容创作、互动优化等一站式服务。')
quote_run.italic = True

doc.add_page_break()

# ========== 二、需求背景 ==========
add_heading(doc, '二、需求背景', 1)
add_heading(doc, '2.1 市场痛点', 2)

table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'

headers2 = ['痛点', '描述', '影响']
data2 = [
    ['内容定位模糊', '不知道自己适合做什么方向，难以形成独特风格', '内容散乱，粉丝粘性低'],
    ['选题低效', '花费大量时间思考拍什么，选题靠碰运气', '创作效率低，更新不稳定'],
    ['运营不专业', '不懂平台算法、发布时间、互动技巧', '流量获取困难'],
    ['涨粉困难', '缺乏系统性的成长路径和策略', '账号发展缓慢'],
]

for i, header in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data2):
    for col_idx, text in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = text

doc.add_paragraph()
add_heading(doc, '2.2 目标用户', 2)
add_para(doc, '初级KOC：粉丝数 < 1万，内容创作经验不足', bold=True)
add_para(doc, '转型KOC：从其他行业转向社交媒体，需要快速上手', bold=True)
add_para(doc, '副业KOC：利用业余时间运营账号，需要高效工具', bold=True)

doc.add_page_break()

# ========== 三、产品介绍 ==========
add_heading(doc, '三、产品介绍', 1)
add_heading(doc, '3.1 产品定位', 2)
add_para(doc, 'KOC AI增长助手是一款面向普通KOC的社媒AI Agent，通过AI能力提供内容方向、选题建议、发布策略、互动优化、AI小游戏等服务，帮助其低成本、高效率实现账号成长。')

doc.add_paragraph()
add_heading(doc, '3.2 核心价值主张', 2)

quote2 = doc.add_paragraph()
quote2_run = quote2.add_run('"让每个普通人都能轻松玩转社交媒体"')
quote2_run.bold = True
quote2_run.italic = True

add_para(doc, 'AI驱动：集成大语言模型，提供智能化服务', bold=True)
add_para(doc, '一站式：覆盖账号运营全流程', bold=True)
add_para(doc, '低成本：无需专业团队，个人即可使用', bold=True)
add_para(doc, '高效率：AI辅助决策，大幅提升运营效率', bold=True)

doc.add_page_break()

# ========== 四、核心功能 ==========
add_heading(doc, '四、核心功能', 1)
add_heading(doc, '4.1 功能架构', 2)
add_para(doc, 'KOC AI增长助手包含六大核心模块：AI账号诊断、智能选题引擎、内容创作助手、AI小游戏生成、数据看板、成长路径规划。')

doc.add_paragraph()
add_heading(doc, '4.2 功能详解', 2)
add_heading(doc, '4.2.1 AI账号诊断', 3)
add_para(doc, '用户输入账号名称和简介后，AI深度分析账号定位与内容策略，输出内容包括：')
add_para(doc, '1. 账号定位分析 - 当前定位评估、优势分析、改进建议')
add_para(doc, '2. 内容策略规划 - 内容矩阵建议、发布频率、最佳发布时间')
add_para(doc, '3. 受众画像分析 - 核心受众特征、受众需求洞察、兴趣标签')
add_para(doc, '4. 30天涨粉计划 - 每周具体行动项、关键里程碑、目标指标')
add_para(doc, '5. 账号健康度评分 - 综合评分（0-100）、各维度雷达图')

add_heading(doc, '4.2.2 智能选题引擎', 3)
add_para(doc, '基于内容赛道和创意程度，AI生成个性化选题建议。')
add_para(doc, '输入参数：内容赛道（8大分类）、创意程度（3级）、选题数量（3-10个）', bold=True)
add_para(doc, '输出内容：选题标题+描述、难度等级、相关标签、今日热点结合', bold=True)

add_heading(doc, '4.2.3 内容创作助手', 3)

table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Table Grid'

headers3 = ['工具', '输入', '输出']
data3 = [
    ['视频脚本生成', '视频主题、时长、风格', '完整分镜脚本（开头吸睛→核心价值→互动引导→结尾升华）'],
    ['标题优化', '内容主题、目标平台', '5-10个备选标题 + 吸引力评分'],
    ['文案生成', '内容主题', '完整社交媒体文案'],
    ['封面建议', '视频主题', '封面设计建议（视觉/文字/色彩/元素）'],
]

for i, header in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data3):
    for col_idx, text in enumerate(row_data):
        table3.rows[row_idx + 1].cells[col_idx].text = text

add_heading(doc, '4.2.4 AI小游戏生成（赛道核心功能）', 3)
add_para(doc, '基于游戏主题和类型，AI一键生成完整的小游戏设计方案，包含规则、分享机制、视觉素材建议。')
add_para(doc, '输入参数：', bold=True)
add_para(doc, '游戏主题：用户自定义（如：测测你的反应速度）')
add_para(doc, '游戏类型：挑战类、知识问答类、互动测试类、抽奖类')
add_para(doc, '时长要求：5-30秒')
add_para(doc, '风格：欢乐有趣、酷炫潮流、可爱清新')

doc.add_page_break()

# ========== 五、技术架构 ==========
add_heading(doc, '五、技术架构', 1)
add_heading(doc, '5.1 技术栈', 2)

table4 = doc.add_table(rows=6, cols=3)
table4.style = 'Table Grid'

headers4 = ['层级', '技术选型', '说明']
data4 = [
    ['前端框架', 'HTML5 + CSS3 + JS', '单页应用，无框架依赖'],
    ['后端框架', 'FastAPI', '高性能异步API框架'],
    ['Agent框架', 'LangGraph', '状态机驱动的Agent编排'],
    ['AI集成', 'OpenAI兼容API', '支持混元、DeepSeek、Claude等'],
    ['数据验证', 'Pydantic', '请求/响应模型定义'],
]

for i, header in enumerate(headers4):
    cell = table4.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data4):
    for col_idx, text in enumerate(row_data):
        table4.rows[row_idx + 1].cells[col_idx].text = text

doc.add_paragraph()
add_heading(doc, '5.2 核心API端点', 2)

table5 = doc.add_table(rows=6, cols=3)
table5.style = 'Table Grid'

headers5 = ['端点', '方法', '说明']
data5 = [
    ['/api/diagnosis', 'POST', 'AI账号诊断'],
    ['/api/topics', 'POST', '智能选题生成'],
    ['/api/content', 'POST', '内容创作（脚本/标题/文案/封面）'],
    ['/api/game', 'POST', 'AI小游戏生成'],
    ['/api/test', 'POST', 'API连接测试'],
]

for i, header in enumerate(headers5):
    cell = table5.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data5):
    for col_idx, text in enumerate(row_data):
        table5.rows[row_idx + 1].cells[col_idx].text = text

doc.add_page_break()

# ========== 六、产品亮点 ==========
add_heading(doc, '六、产品亮点', 1)
add_heading(doc, '6.1 核心创新点', 2)

table6 = doc.add_table(rows=6, cols=2)
table6.style = 'Table Grid'

headers6 = ['亮点', '说明']
data6 = [
    ['开放API架构', '支持用户使用自己的API Key，接入任意OpenAI兼容模型'],
    ['LangGraph Agent', '状态机驱动的Agent编排，可扩展性强'],
    ['纯前端实现', '无需复杂后端，降低部署成本'],
    ['一站式服务', '覆盖账号运营全流程，无需切换多工具'],
    ['AI小游戏集成', '结合赛道方向，支持一键生成互动小游戏'],
]

for i, header in enumerate(headers6):
    cell = table6.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data6):
    for col_idx, text in enumerate(row_data):
        table6.rows[row_idx + 1].cells[col_idx].text = text

add_heading(doc, '6.2 与PCG业务结合', 2)
add_para(doc, '可集成视频号API，直接发布内容')
add_para(doc, '可接入腾讯广告智能投放平台')
add_para(doc, '可结合AIGC能力自动生成广告素材')
add_para(doc, '可结合AI小游戏实现用户增长和品牌传播', bold=True)

doc.add_page_break()

# ========== 七、商业模式 ==========
add_heading(doc, '七、商业模式', 1)
add_heading(doc, '7.1 定价策略', 2)

table7 = doc.add_table(rows=4, cols=3)
table7.style = 'Table Grid'

headers7 = ['版本', '价格', '功能权益']
data7 = [
    ['Free', '免费', '基础API配置，每周10次诊断，20个选题，基础看板'],
    ['Pro', '¥99/月', '无限AI调用，高级数据分析，多账号管理，AI小游戏无限生成'],
    ['Enterprise', '定制', '私有化部署，自定义模型，API批量调用'],
]

for i, header in enumerate(headers7):
    cell = table7.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data7):
    for col_idx, text in enumerate(row_data):
        table7.rows[row_idx + 1].cells[col_idx].text = text

add_heading(doc, '7.2 变现路径', 2)
add_para(doc, '1. 订阅制：月费/年费订阅')
add_para(doc, '2. 增值服务：高级功能单独收费（如AI小游戏模板）')
add_para(doc, '3. 企业版：为企业客户提供定制服务')
add_para(doc, '4. 生态合作：与MCN机构、品牌合作')

doc.add_page_break()

# ========== 八、未来规划 ==========
add_heading(doc, '八、未来规划', 1)

table8 = doc.add_table(rows=5, cols=3)
table8.style = 'Table Grid'

headers8 = ['阶段', '时间', '目标']
data8 = [
    ['Phase 1', '当前', '完成Demo版，验证核心功能（诊断/选题/创作/游戏）'],
    ['Phase 2', '1-2月', '接入真实数据源（视频号API），优化AI模型'],
    ['Phase 3', '3-6月', '推出AI小游戏编辑器，支持自定义规则和品牌元素'],
    ['Phase 4', '6-12月', '商业化探索，接入品牌合作，建立KOC增长生态'],
]

for i, header in enumerate(headers8):
    cell = table8.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data8):
    for col_idx, text in enumerate(row_data):
        table8.rows[row_idx + 1].cells[col_idx].text = text

doc.add_page_break()

# ========== 九、参赛材料清单 ==========
add_heading(doc, '九、参赛材料清单', 1)

table9 = doc.add_table(rows=4, cols=3)
table9.style = 'Table Grid'

headers9 = ['材料', '说明', '状态']
data9 = [
    ['产品Demo', 'Web端交互演示（前后端联调版本）', '已完成'],
    ['产品录屏', '3分钟功能演示视频', '待录制'],
    ['说明文档', '本文档', '已完成'],
]

for i, header in enumerate(headers9):
    cell = table9.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data9):
    for col_idx, text in enumerate(row_data):
        table9.rows[row_idx + 1].cells[col_idx].text = text

doc.add_page_break()

# ========== 十、团队信息 ==========
add_heading(doc, '十、团队信息', 1)

table10 = doc.add_table(rows=5, cols=2)
table10.style = 'Table Grid'

headers10 = ['项目', '内容']
data10 = [
    ['参赛赛道', '赛道⑤ - PCG技术线-智能营销/AI小游戏方向'],
    ['作品名称', 'KOC AI增长助手'],
    ['技术栈', '前端：HTML5+CSS3+JS；后端：FastAPI+LangGraph'],
    ['Demo访问', '本地：http://localhost:8080（前端）+ http://localhost:8000（后端API）'],
]

for i, header in enumerate(headers10):
    cell = table10.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, 'D5E8F0')
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data10):
    for col_idx, text in enumerate(row_data):
        table10.rows[row_idx + 1].cells[col_idx].text = text

# ========== 结尾 ==========
doc.add_paragraph()
doc.add_paragraph()
end = doc.add_paragraph('— 文档结束 —')
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in end.runs:
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(10)

# 保存文档
output_path = 'D:/WorkBuddy/koc-ai-demo/KOC_AI增长助手_说明文档.docx'
doc.save(output_path)
print(f'文档生成成功！保存至：{output_path}')
